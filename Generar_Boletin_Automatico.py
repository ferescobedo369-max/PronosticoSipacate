"""
Generador Automático de Boletines de Viento - Sipacate
Genera boletines en formato Word con pronósticos de viento y análisis automático
Versión 2.0 - Con análisis inteligente automático
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
import locale

# Importar módulo de análisis
from analisis_viento import (
    leer_datos_pronostico,
    generar_interpretacion_automatica,
    generar_interpretacion_velocidad,
    validar_datos,
    obtener_estadisticas_dia
)

# Configurar locale a español (si está disponible)
try:
    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_TIME, 'Spanish_Spain.1252')
    except:
        print("⚠️ No se pudo configurar locale español, usando inglés")

# ===================== CONFIGURACIÓN DE RUTAS =====================
RUTA_BASE = r"C:\DATA\OneDrive - Asazgua\SSP\Información generada por Fernando\Zafra 25-26\Sipacate"
PATH_PLANTILLA = os.path.join(RUTA_BASE, "Boletines", "plantilla de boletin.docx")
PATH_IMAGENES = os.path.join(RUTA_BASE, "Pronosticos", "Wind_forecasts", "2_Results")
PATH_SALIDA = os.path.join(RUTA_BASE, "Boletines")

# ===================== FUNCIONES AUXILIARES =====================

def obtener_nombre_dia_espanol(fecha):
    """Convierte fecha a nombre del día en español"""
    dias = {
        'Monday': 'lunes',
        'Tuesday': 'martes',  
        'Wednesday': 'miércoles',
        'Thursday': 'jueves',
        'Friday': 'viernes',
        'Saturday': 'sábado',
        'Sunday': 'domingo'
    }
    nombre_ingles = fecha.strftime('%A')
    return dias.get(nombre_ingles, nombre_ingles)

def obtener_nombre_mes_espanol(fecha):
    """Convierte fecha a nombre del mes en español"""
    meses = {
        'January': 'enero',
        'February': 'febrero',
        'March': 'marzo',
        'April': 'abril',
        'May': 'mayo',
        'June': 'junio',
        'July': 'julio',
        'August': 'agosto',
        'September': 'septiembre',
        'October': 'octubre',
        'November': 'noviembre',
        'December': 'diciembre'
    }
    nombre_ingles = fecha.strftime('%B')
    return meses.get(nombre_ingles, nombre_ingles)

def buscar_graficas_pronostico(fecha, path_imagenes):
    """
    Busca las 4 gráficas de serie de tiempo para el boletín
    Ahora busca en carpeta raíz de 2_Results
    """
    fecha_str = fecha.strftime("%Y%m%d")
    graficas = {
        'direccion_10m': None,
        'direccion_100m': None,
        'velocidad_10m': None,
        'velocidad_100m': None
    }
    
    for archivo in os.listdir(path_imagenes):
        if fecha_str in archivo and archivo.endswith(".png"):
            if "Serie_ciclo_diario_wind_direction_10m" in archivo:
                graficas['direccion_10m'] = os.path.join(path_imagenes, archivo)
            elif "Serie_ciclo_diario_wind_direction_100m" in archivo:
                graficas['direccion_100m'] = os.path.join(path_imagenes, archivo)
            elif "Serie_ciclo_diario_wind_speed_10m" in archivo:
                graficas['velocidad_10m'] = os.path.join(path_imagenes, archivo)
            elif "Serie_ciclo_diario_wind_speed_100m" in archivo:
                graficas['velocidad_100m'] = os.path.join(path_imagenes, archivo)
    
    return graficas

# ===================== FUNCIÓN PRINCIPAL =====================

def generar_boletin_automatico(fecha_inicio=None):
    """
    Genera un boletín de viento en formato Word con análisis automático
    
    Parámetros:
    - fecha_inicio: datetime object. Si es None, usa la fecha actual
    """
    
    # Usar fecha actual si no se especifica
    if fecha_inicio is None:
        fecha_inicio = datetime.now()
    
    # Generar correlativo (BO + DDMMYY)
    correlativo = f"BO{fecha_inicio.strftime('%d%m%y')}"
    
    print(f"\n{'='*60}")
    print(f"🔄 Generando Boletín {correlativo}")
    print(f"{'='*60}")
    
    # Verificar que existe la plantilla
    if not os.path.exists(PATH_PLANTILLA):
        print(f"❌ Error: No se encontró la plantilla en:\n   {PATH_PLANTILLA}")
        return False
    
    if not os.path.exists(PATH_IMAGENES):
        print(f"❌ Error: No se encontró la carpeta de imágenes en:\n   {PATH_IMAGENES}")
        return False
    
    # ============ LEER Y VALIDAR DATOS DE PRONÓSTICO ============
    print(f"\n📊 Cargando datos de pronóstico...")
    df_pronostico = leer_datos_pronostico(fecha_inicio, os.path.join(RUTA_BASE, "Pronosticos", "Wind_forecasts"))
    
    if not validar_datos(df_pronostico):
        print(f"❌ Error: Datos de pronóstico no válidos")
        return False
    
    # Cargar plantilla
    print(f"\n📄 Cargando plantilla...")
    doc = Document(PATH_PLANTILLA)
    
    # ============ ENCABEZADO DEL BOLETÍN ============
    print(f"✍️  Agregando encabezado...")
    
    # Correlativo
    p_correlativo = doc.add_paragraph()
    p_correlativo.add_run(f"Correlativo {correlativo}").bold = True
    p_correlativo.paragraph_format.space_after = Pt(6)
    
    # ICC
    p_icc = doc.add_paragraph("ICC: CeH y SSP")
    p_icc.paragraph_format.space_after = Pt(12)
    
    # Título principal
    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("Condiciones de viento proyectadas para las zonas buffer prioritaria")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)
    p_titulo.paragraph_format.space_after = Pt(12)
    
    # Fechas del pronóstico (3 días consecutivos)
    dia1 = fecha_inicio
    dia2 = fecha_inicio + timedelta(days=1)
    dia3 = fecha_inicio + timedelta(days=2)
    
    mes = obtener_nombre_mes_espanol(dia1)
    anio = dia1.year
    
    texto_fechas = f"Pronósticos de horarios indicados para el {dia1.day}, {dia2.day} de {obtener_nombre_mes_espanol(dia1)} y {dia3.day} de {obtener_nombre_mes_espanol(dia3)} {anio}."
    p_fechas = doc.add_paragraph(texto_fechas)
    p_fechas.paragraph_format.space_after = Pt(12)
    
    # ============ INTERPRETACIÓN GENERAL (CONSTANTE) ============
    print(f"📋 Agregando texto de interpretación general...")
    
    p_interp_titulo = doc.add_paragraph()
    p_interp_titulo.add_run("Interpretación: ").bold = True
    p_interp_titulo.add_run("""Para facilitar el análisis, la gráfica utiliza un sistema de colores: las áreas en rosado representan condiciones no favorables (no aptas para realizar la quema), mientras que las franjas en verde indican periodos favorables para la operación.
Cada línea de color en el gráfico corresponde a una fecha específica del calendario. En la línea horizontal, se detallan las horas del día (de 0 a 23:59), donde cada cuadrícula representa un intervalo de 2 horas.""")
    p_interp_titulo.paragraph_format.space_after = Pt(6)
    
    # Zonas no favorables
    p_zona_rosa = doc.add_paragraph()
    p_zona_rosa.add_run("• Zonas no favorables (franjas rosadas): ").bold = True
    p_zona_rosa.add_run("Representan los horarios donde el viento sopla predominantemente hacia el Sur (entre los 0° y 90°, y de 270° a 360°).")
    
    # Zonas favorables
    p_zona_verde = doc.add_paragraph()
    p_zona_verde.add_run("• Zonas favorables (franja verde): ").bold = True
    p_zona_verde.add_run("Corresponden a los momentos en que el viento mantiene una dirección hacia el Norte (entre los 90° y 270°). Estas condiciones son las adecuadas para realizar la actividad de quemas ya que se espera un desplazamiento de la pavesa al norte.")
    
    # Líneas de colores
    p_lineas = doc.add_paragraph()
    p_lineas.add_run("Líneas de colores: ").bold = True
    p_lineas.add_run("Significan una fecha en el calendario.")
    p_lineas.paragraph_format.space_after = Pt(12)
    
    # ============ INSERTAR GRÁFICAS ============
    print(f"\n🖼️  Buscando e insertando gráficas...")
    
    graficas = buscar_graficas_pronostico(fecha_inicio, PATH_IMAGENES)
    
    # Insertar las 4 gráficas en orden
    graficas_orden = [
        ('direccion_10m', 'Dirección del viento a 10 metros'),
        ('direccion_100m', 'Dirección del viento a 100 metros'),
        ('velocidad_10m', 'Velocidad del viento a 10 metros'), 
        ('velocidad_100m', 'Velocidad del viento a 100 metros')
    ]
    
    for key, titulo_base in graficas_orden:
        if graficas[key]:
            print(f"   ✅ Insertando gráfica: {key}")
            
            # TÍTULO EN NEGRITA ARRIBA DE LA GRÁFICA
            p = doc.add_paragraph()
            run = p.add_run(titulo_base)
            run.bold = True
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            # IMAGEN
            doc.add_picture(graficas[key], width=Inches(6.0))
            doc.add_paragraph().paragraph_format.space_after = Pt(12)
        else:
            print(f"   ⚠️ No se encontró gráfica: {key}")
    
    # ============ ANÁLISIS POR DÍA (AUTOMÁTICO) ============
    print(f"\n📊 Generando análisis automático de viento por día...")
    
    # Título de sección
    p_analisis_titulo = doc.add_paragraph()
    run_analisis = p_analisis_titulo.add_run("Dirección del viento")
    run_analisis.bold = True
    run_analisis.font.size = Pt(13)
    p_analisis_titulo.paragraph_format.space_before = Pt(12)
    p_analisis_titulo.paragraph_format.space_after = Pt(12)
    
    fechas_analisis = [dia1, dia2, dia3]
    
    for idx, fecha in enumerate(fechas_analisis):
        dia_nombre = obtener_nombre_dia_espanol(fecha)
        dia_num = fecha.day
        mes_nombre = obtener_nombre_mes_espanol(fecha)
        
        # Filtrar datos del día específico
        df_dia = df_pronostico[df_pronostico['date'].dt.date == fecha.date()].copy()
        
        if df_dia.empty:
            print(f"   ⚠️ No hay datos para {dia_nombre} {dia_num}")
            continue
        
        # Obtener estadísticas para debugging
        stats = obtener_estadisticas_dia(fecha, df_dia)
        print(f"   📈 {dia_nombre} {dia_num}: {stats['10m']['horas_favorables']}h favorables, "
              f"{stats['10m']['transiciones']} transiciones")
        
        # GENERAR INTERPRETACIÓN AUTOMÁTICA
        interpretacion = generar_interpretacion_automatica(
            fecha,
            df_dia,
            dia_nombre,
            dia_num,
            mes_nombre
        )
        
        # Agregar al documento
        p = doc.add_paragraph()
        p.add_run(f"{dia_nombre.capitalize()} {dia_num} de {mes_nombre}: ").bold = True
        p.add_run(interpretacion)
        p.paragraph_format.space_after = Pt(12)
        
        print(f"   ✅ Análisis generado para {dia_nombre} {dia_num}")
    
    # ============ ANÁLISIS DE VELOCIDAD POR DÍA (AUTOMÁTICO) ============
    print(f"\n🌬️  Generando análisis automático de velocidad por día...")
    
    # Título de sección de velocidad
    p_velocidad_titulo = doc.add_paragraph()
    run_velocidad = p_velocidad_titulo.add_run("Velocidad del viento")
    run_velocidad.bold = True
    run_velocidad.font.size = Pt(13)
    p_velocidad_titulo.paragraph_format.space_before = Pt(12)
    p_velocidad_titulo.paragraph_format.space_after = Pt(12)
    
    for idx, fecha in enumerate(fechas_analisis):
        dia_nombre = obtener_nombre_dia_espanol(fecha)
        dia_num = fecha.day
        mes_nombre = obtener_nombre_mes_espanol(fecha)
        
        # Filtrar datos del día específico
        df_dia = df_pronostico[df_pronostico['date'].dt.date == fecha.date()].copy()
        
        if df_dia.empty:
            print(f"   ⚠️ No hay datos para {dia_nombre} {dia_num}")
            continue
        
        # GENERAR INTERPRETACIÓN AUTOMÁTICA DE VELOCIDAD
        interpretacion_vel = generar_interpretacion_velocidad(
            fecha,
            df_dia,
            dia_nombre,
            dia_num,
            mes_nombre
        )
        
        # Agregar al documento
        p = doc.add_paragraph()
        p.add_run(f"{dia_nombre.capitalize()} {dia_num} de {mes_nombre}: ").bold = True
        p.add_run(interpretacion_vel)
        p.paragraph_format.space_after = Pt(12)
        
        print(f"   ✅ Análisis de velocidad generado para {dia_nombre} {dia_num}")
    
    # ============ GUARDAR DOCUMENTO ============
    os.makedirs(PATH_SALIDA, exist_ok=True)
    
    nombre_archivo = f"Boletin_Sipacate_{correlativo}.docx"
    ruta_completa = os.path.join(PATH_SALIDA, nombre_archivo)
    
    doc.save(ruta_completa)
    
    print(f"\n{'='*60}")
    print(f"✅ BOLETÍN GENERADO EXITOSAMENTE (Con análisis automático)")
    print(f"{'='*60}")
    print(f"📁 Ubicación: {ruta_completa}")
    print(f"📄 Nombre: {nombre_archivo}")
    print(f"📅 Fecha: {fecha_inicio.strftime('%d/%m/%Y')}")
    print(f"🔖 Correlativo: {correlativo}")
    print(f"🤖 Análisis: Generado automáticamente de datos CSV")
    print(f"{'='*60}\n")
    
    return True

# ===================== EJECUCIÓN =====================

if __name__ == "__main__":
    print("""
    ╔════════════════════════════════════════════════════════════╗
    ║   GENERADOR AUTOMÁTICO DE BOLETINES DE VIENTO - SIPACATE  ║
    ║              Con Análisis Inteligente v2.0                 ║
    ╚════════════════════════════════════════════════════════════╝
    """)
    
    # Generar boletín con fecha actual
    exito = generar_boletin_automatico()
    
    if exito:
        print("\n💡 El boletín se generó con análisis automático de los datos.")
        print("💡 Las interpretaciones se crearon analizando las direcciones de viento.")
        print("💡 Puedes abrir el archivo Word para revisarlo y agregar recomendaciones.")
    else:
        print("\n❌ Hubo un error al generar el boletín. Verifica las rutas y archivos.")
    
    input("\nPresiona ENTER para salir...")
