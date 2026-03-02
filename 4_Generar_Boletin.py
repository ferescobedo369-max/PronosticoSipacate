"""
Script 4: Generador Automático de Boletín de Viento - Sipacate
- Lee Area_forecast_latest.csv generado por Script 1
- Genera análisis automático de dirección y velocidad
- Produce Boletin_Sipacate_BODDMMYY.docx (usando plantilla con logo ICC)
- Convierte a PDF con LibreOffice (disponible en GitHub Actions)
- Guarda PDF en carpeta YYYYMMDD y como Boletin_latest.pdf
"""

import os
import shutil
import subprocess
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# ---------------------------------------------------------
# Rutas
# ---------------------------------------------------------
base_dir    = os.path.dirname(os.path.abspath(__file__))
results_dir = os.path.join(base_dir, "2_Results")
# Buscar plantilla en ubicaciones posibles
def _encontrar_plantilla(base_dir):
    candidatos = [
        os.path.join(base_dir, "0_Archivos_Recibidos", "plantilla_de_boletin.docx"),
        os.path.join(base_dir, "plantilla_de_boletin.docx"),
    ]
    # Búsqueda recursiva por si está en subcarpeta
    for root, dirs, files in os.walk(base_dir):
        for f in files:
            if f.lower() == "plantilla_de_boletin.docx":
                candidatos.insert(0, os.path.join(root, f))
    for c in candidatos:
        if os.path.exists(c):
            print(f"✅ Plantilla encontrada en: {c}")
            return c
    raise FileNotFoundError(
        f"No se encontró plantilla_de_boletin.docx en el repo.\n"
        f"Rutas buscadas:\n" + "\n".join(candidatos)
    )

plantilla = _encontrar_plantilla(base_dir)
csv_path    = os.path.join(results_dir, "Area_forecast_latest.csv")

# ---------------------------------------------------------
# Helpers fechas en español
# ---------------------------------------------------------
DIAS_ES = {
    'Monday':'lunes','Tuesday':'martes','Wednesday':'miércoles',
    'Thursday':'jueves','Friday':'viernes','Saturday':'sábado','Sunday':'domingo'
}
MESES_ES = {
    'January':'enero','February':'febrero','March':'marzo','April':'abril',
    'May':'mayo','June':'junio','July':'julio','August':'agosto',
    'September':'septiembre','October':'octubre','November':'noviembre','December':'diciembre'
}

def dia_es(fecha):
    return DIAS_ES.get(fecha.strftime('%A'), fecha.strftime('%A'))

def mes_es(fecha):
    return MESES_ES.get(fecha.strftime('%B'), fecha.strftime('%B'))

# ---------------------------------------------------------
# Leer CSV
# ---------------------------------------------------------
def leer_datos(csv_path):
    df = pd.read_csv(csv_path)
    df['date'] = pd.to_datetime(df['date'], utc=True).dt.tz_convert('America/Guatemala')
    df['date'] = df['date'].dt.tz_localize(None)
    df['Hora'] = df['date'].dt.hour
    return df

# ---------------------------------------------------------
# Análisis DIRECCIÓN
# ---------------------------------------------------------
def analizar_direccion(df_dia):
    dir_10m = df_dia['wind_direction_10m'].values
    horas   = df_dia['Hora'].values

    def es_favorable(g):
        return 90 <= g <= 270

    favorable_mask = np.array([es_favorable(d) for d in dir_10m])
    horas_nofav    = horas[~favorable_mask]

    # Bloque favorable principal
    max_bloque, bloque_actual = 0, 0
    inicio_fav, fin_fav, inicio_actual = None, None, None

    for i, (h, fav) in enumerate(zip(horas, favorable_mask)):
        if fav:
            if bloque_actual == 0:
                inicio_actual = h
            bloque_actual += 1
        else:
            if bloque_actual > max_bloque:
                max_bloque = bloque_actual
                inicio_fav = inicio_actual
                fin_fav    = horas[i - 1] if i > 0 else h
            bloque_actual = 0

    if bloque_actual > max_bloque:
        max_bloque = bloque_actual
        inicio_fav = inicio_actual
        fin_fav    = horas[-1]

    madrugada_nofav = int(sum(1 for h, fav in zip(horas, favorable_mask) if h < 7 and not fav))
    noche_nofav     = int(sum(1 for h, fav in zip(horas, favorable_mask) if h >= 20 and not fav))

    partes = []

    if madrugada_nofav >= 4:
        hora_limite = int(horas_nofav[horas_nofav < 10].max()) if len(horas_nofav[horas_nofav < 10]) else 7
        partes.append(
            f"Se identifica una alta variabilidad del viento durante la madrugada y las "
            f"primeras horas de la mañana (00:00 - {hora_limite:02d}:00), con el viento soplando "
            f"frecuentemente en la zona no favorable, es decir hacia el sur."
        )
    else:
        partes.append(
            "Durante las primeras horas del día se presentan condiciones variables, "
            "con algunas horas mostrando viento hacia zonas no favorables."
        )

    if inicio_fav is not None and max_bloque >= 4:
        partes.append(
            f"A partir de las {inicio_fav:02d}:00 horas, la dirección se consolida dentro de la "
            f"zona favorable (90° a 270°) en ambos niveles de altura (10 y 100 metros), "
            f"manteniendo estabilidad durante la tarde hasta las {min(int(fin_fav)+1, 23):02d}:00 horas aproximadamente."
        )
    elif int(favorable_mask.sum()) >= 6:
        partes.append(
            "La dirección se mantiene dentro de la zona favorable durante buena parte del día, "
            "especialmente durante las horas de la tarde."
        )
    else:
        partes.append(
            "El viento presenta condiciones predominantemente no favorables durante la mayor "
            "parte del día, con pocas horas dentro de la zona verde."
        )

    if noche_nofav >= 2:
        hora_cambio = next((int(h) for h, fav in zip(horas, favorable_mask) if h >= 19 and not fav), 20)
        partes.append(
            f"El comportamiento vuelve a mostrar variaciones hacia la zona roja después de "
            f"las {hora_cambio:02d}:00 horas."
        )

    return " ".join(partes)

# ---------------------------------------------------------
# Análisis VELOCIDAD
# ---------------------------------------------------------
def analizar_velocidad(df_dia):
    v10   = df_dia['wind_speed_10m'].values
    v100  = df_dia['wind_speed_100m'].values
    horas = df_dia['Hora'].values

    prom_10  = float(np.mean(v10))
    prom_100 = float(np.mean(v100))
    max_10   = float(np.max(v10))
    hora_max = int(horas[np.argmax(v10)])

    sobre_umbral = v10 >= 10
    n_sobre      = int(sobre_umbral.sum())
    if n_sobre > 0:
        h_ini = int(horas[sobre_umbral][0])
        h_fin = int(horas[sobre_umbral][-1])

    v_manana = float(np.mean(v10[horas < 10]))  if any(horas < 10)  else 0.0
    v_tarde  = float(np.mean(v10[(horas >= 12) & (horas <= 19)])) if any((horas >= 12) & (horas <= 19)) else 0.0
    diff     = prom_100 - prom_10

    intensidad = "moderado a fuerte" if max_10 >= 20 else "moderado" if max_10 >= 10 else "leve a moderado"

    partes = [
        f"Las velocidades promedio del viento se registran en {prom_10:.1f} km/h a 10 metros "
        f"de altura y {prom_100:.1f} km/h a 100 metros.",
        f"Las velocidades máximas alcanzan {max_10:.1f} km/h alrededor de las {hora_max:02d}:00 horas "
        f"a 10 metros de altura, indicando condiciones de viento {intensidad} durante este periodo."
    ]

    if n_sobre > 0:
        partes.append(
            f"Se pronostican velocidades por encima de 10 km/h durante {n_sobre} horas del día, "
            f"principalmente entre las {h_ini:02d}:00 y {h_fin:02d}:00 horas."
        )
    else:
        partes.append("Las velocidades se mantienen por debajo del umbral de 10 km/h durante todo el día.")

    if v_tarde > v_manana and v_tarde > 5:
        partes.append(
            f"Se observa un incremento significativo de las velocidades durante la tarde, "
            f"pasando de {v_manana:.1f} km/h en la mañana a {v_tarde:.1f} km/h en la tarde."
        )

    if diff > 4:
        partes.append(
            f"Se observa una diferencia notable entre niveles, con velocidades {diff:.1f} km/h "
            f"más altas a 100 metros de altura, lo cual es característico del perfil vertical del viento."
        )

    return " ".join(partes)

# ---------------------------------------------------------
# CONVERTIR DOCX → PDF con LibreOffice
# ---------------------------------------------------------
def convertir_a_pdf(docx_path, output_dir):
    """
    Usa LibreOffice headless para convertir .docx a .pdf
    Disponible por defecto en ubuntu-latest de GitHub Actions
    """
    print("🔄 Convirtiendo Word → PDF con LibreOffice...")
    resultado = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf",
         "--outdir", output_dir, docx_path],
        capture_output=True, text=True, timeout=120
    )
    if resultado.returncode != 0:
        print(f"⚠️  stdout: {resultado.stdout}")
        print(f"⚠️  stderr: {resultado.stderr}")
        raise RuntimeError(f"LibreOffice falló al convertir: {resultado.stderr}")

    # LibreOffice nombra el PDF igual que el docx pero con .pdf
    nombre_pdf = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
    pdf_path   = os.path.join(output_dir, nombre_pdf)

    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"No se generó el PDF en: {pdf_path}")

    print(f"✅ PDF generado: {pdf_path}")
    return pdf_path

# ---------------------------------------------------------
# FUNCIÓN PRINCIPAL
# ---------------------------------------------------------
def generar_boletin():

    if not os.path.exists(csv_path):
        raise FileNotFoundError(f"No se encontró CSV de pronóstico en: {csv_path}")

    print("📊 Leyendo datos...")
    df = leer_datos(csv_path)
    df['date_only'] = df['date'].dt.date

    fechas_unicas = sorted(df['date_only'].unique())
    n_dias  = min(3, len(fechas_unicas))
    fechas  = [datetime.combine(d, datetime.min.time()) for d in fechas_unicas[:n_dias]]

    if not fechas:
        raise RuntimeError("No hay datos disponibles.")

    fecha_inicio = fechas[0]
    correlativo  = f"BO{fecha_inicio.strftime('%d%m%y')}"
    run_date_str = fecha_inicio.strftime("%Y%m%d")
    carpeta_fecha = os.path.join(results_dir, run_date_str)
    os.makedirs(carpeta_fecha, exist_ok=True)

    print(f"📅 Generando boletín {correlativo}...")

    # ---- Abrir plantilla ----
    doc = Document(plantilla)
    for p in doc.paragraphs:
        p.clear()

    def add_para(bold=None, normal=None, size_pt=None, space_after=6, space_before=0,
                 justify=True):
        p = doc.add_paragraph()
        if bold:
            r = p.add_run(bold)
            r.bold = True
            if size_pt:
                r.font.size = Pt(size_pt)
        if normal:
            r2 = p.add_run(normal)
            r2.bold = False
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        if justify:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    # ---- Encabezado ----
    add_para(bold=f"Correlativo {correlativo}", space_after=4)
    add_para(bold="ICC: CeH y SSP", space_after=4)
    add_para(bold="Condiciones de viento proyectadas para las zonas buffer prioritaria",
             size_pt=14, space_after=6)

    dia1 = fechas[0]
    dia2 = fechas[1] if n_dias > 1 else fechas[0]
    dia3 = fechas[2] if n_dias > 2 else fechas[0]

    if n_dias == 3:
        texto_fechas = (
            f"Pronósticos de horarios indicados para el {dia1.day}, {dia2.day} de {mes_es(dia1)} "
            f"y {dia3.day} de {mes_es(dia3)} {dia1.year}."
        )
    elif n_dias == 2:
        texto_fechas = (
            f"Pronósticos de horarios indicados para el {dia1.day} y {dia2.day} de "
            f"{mes_es(dia1)} {dia1.year}."
        )
    else:
        texto_fechas = f"Pronóstico para el {dia1.day} de {mes_es(dia1)} {dia1.year}."

    add_para(normal=texto_fechas, space_after=10)

    # ---- Interpretación general ----
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Interpretación: ").bold = True
    p.add_run(
        "Para facilitar el análisis, la gráfica utiliza un sistema de colores: las áreas en rosado "
        "representan condiciones no favorables (no aptas para realizar la quema), mientras que las "
        "franjas en verde indican periodos favorables para la operación.\n"
        "Cada línea de color en el gráfico corresponde a una fecha específica del calendario. "
        "En la línea horizontal, se detallan las horas del día (de 0 a 23:59), donde cada cuadrícula "
        "representa un intervalo de 2 horas."
    )
    p.paragraph_format.space_after = Pt(6)

    p2 = doc.add_paragraph()
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.add_run("• Zonas no favorables (franjas rosadas): ").bold = True
    p2.add_run(
        "Representan los horarios donde el viento sopla predominantemente hacia el Sur "
        "(entre los 0° y 90°, y de 270° a 360°)."
    )
    p2.paragraph_format.space_after = Pt(4)

    p3 = doc.add_paragraph()
    p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.add_run("• Zonas favorables (franja verde): ").bold = True
    p3.add_run(
        "Corresponden a los momentos en que el viento mantiene una dirección hacia el Norte "
        "(entre los 90° y 270°). Estas condiciones son las adecuadas para realizar la actividad "
        "de quemas ya que se espera un desplazamiento de la pavesa al norte."
    )
    p3.paragraph_format.space_after = Pt(4)

    p4 = doc.add_paragraph()
    p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p4.add_run("Líneas de colores: ").bold = True
    p4.add_run("Significan una fecha en el calendario.")
    p4.paragraph_format.space_after = Pt(12)

    # ---- Gráficas ----
    graficas_config = [
        ("Serie_ciclo_diario_wind_direction_10m",  "Dirección del viento a 10 metros"),
        ("Serie_ciclo_diario_wind_direction_100m", "Dirección del viento a 100 metros"),
        ("Serie_ciclo_diario_wind_speed_10m",      "Velocidad del viento a 10 metros"),
        ("Serie_ciclo_diario_wind_speed_100m",     "Velocidad del viento a 100 metros"),
    ]

    carpeta_imgs = carpeta_fecha if os.path.isdir(carpeta_fecha) else results_dir

    print("\n🖼️  Insertando gráficas (2 por página)...")

    # Resolver rutas de imágenes
    rutas_imgs = []
    for prefijo, titulo in graficas_config:
        ruta_img = None
        for archivo in os.listdir(carpeta_imgs):
            if prefijo in archivo and archivo.endswith(".png"):
                ruta_img = os.path.join(carpeta_imgs, archivo)
                break
        rutas_imgs.append((titulo, ruta_img))

    # Insertar de 2 en 2 por página
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def add_page_break(doc):
        p = doc.add_paragraph()
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)
        return p

    for i, (titulo_grafica, ruta_img) in enumerate(rutas_imgs):
        # Título centrado y en negrita
        p_tit = doc.add_paragraph()
        r = p_tit.add_run(titulo_grafica)
        r.bold = True
        r.font.size = Pt(11)
        p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_tit.paragraph_format.space_before = Pt(4)
        p_tit.paragraph_format.space_after  = Pt(4)

        # Imagen
        if ruta_img:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()
            run_img.add_picture(ruta_img, width=Inches(5.8))
            print(f"   ✅ {titulo_grafica}")
        else:
            doc.add_paragraph(f"[Gráfica no disponible: {titulo_grafica}]")
            print(f"   ⚠️  No encontrada: {titulo_grafica}")

        # Cada 2 imágenes (índice impar) insertar salto de página
        # excepto después de la última imagen
        if i % 2 == 1 and i < len(rutas_imgs) - 1:
            add_page_break(doc)
        elif i % 2 == 0 and i < len(rutas_imgs) - 1:
            # Separador entre las dos imágenes de la misma página
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---- Análisis dirección ----
    print("\n📋 Análisis de dirección...")
    p_tit = doc.add_paragraph()
    p_tit.add_run("Dirección del viento").bold = True
    p_tit.runs[0].font.size = Pt(13)
    p_tit.paragraph_format.space_before = Pt(12)
    p_tit.paragraph_format.space_after  = Pt(8)

    for fecha in fechas:
        df_dia = df[df['date_only'] == fecha.date()].copy()
        if df_dia.empty:
            continue
        texto = analizar_direccion(df_dia)
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(f"{dia_es(fecha).capitalize()} {fecha.day} de {mes_es(fecha)}: ").bold = True
        p.add_run(texto)
        p.paragraph_format.space_after = Pt(10)
        print(f"   ✅ {dia_es(fecha)} {fecha.day}")

    # ---- Análisis velocidad ----
    print("\n🌬️  Análisis de velocidad...")
    p_tit2 = doc.add_paragraph()
    p_tit2.add_run("Velocidad del viento").bold = True
    p_tit2.runs[0].font.size = Pt(13)
    p_tit2.paragraph_format.space_before = Pt(12)
    p_tit2.paragraph_format.space_after  = Pt(8)

    for fecha in fechas:
        df_dia = df[df['date_only'] == fecha.date()].copy()
        if df_dia.empty:
            continue
        texto = analizar_velocidad(df_dia)
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(f"{dia_es(fecha).capitalize()} {fecha.day} de {mes_es(fecha)}: ").bold = True
        p.add_run(texto)
        p.paragraph_format.space_after = Pt(10)
        print(f"   ✅ {dia_es(fecha)} {fecha.day}")

    # ---- Guardar DOCX ----
    nombre_docx = f"Boletin_Sipacate_{correlativo}.docx"
    ruta_docx   = os.path.join(carpeta_fecha, nombre_docx)
    doc.save(ruta_docx)
    print(f"\n📄 Word guardado: {ruta_docx}")

    # ---- Convertir a PDF ----
    pdf_path    = convertir_a_pdf(ruta_docx, carpeta_fecha)
    latest_pdf  = os.path.join(results_dir, "Boletin_latest.pdf")
    shutil.copy2(pdf_path, latest_pdf)
    print(f"📎 PDF latest: {latest_pdf}")

    print(f"\n{'='*55}")
    print(f"✅ BOLETÍN COMPLETO: {os.path.basename(pdf_path)}")
    print(f"{'='*55}\n")

    return pdf_path


if __name__ == "__main__":
    generar_boletin()
